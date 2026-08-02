import io
import logging

import docx
from apps.api.core.storage import storage_service
from apps.api.models.draft import Draft
from apps.worker.jobs import TerminalJobError
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from sqlalchemy.ext.asyncio import AsyncSession

logger = logging.getLogger("worker.export")


async def handle_export_draft(payload: dict, session: AsyncSession):
    draft_id = payload.get("draft_id")
    fmt = payload.get("format", "pdf").lower()

    if not draft_id:
        raise TerminalJobError("Missing draft_id in EXPORT_DRAFT payload")

    draft = await session.get(Draft, draft_id)
    if not draft:
        raise TerminalJobError(f"Draft {draft_id} not found for export")

    logger.info(f"Exporting draft {draft_id} (version {draft.version}) as {fmt}")

    content_bytes = b""
    content_type = (
        "application/pdf"
        if fmt == "pdf"
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    if fmt == "docx":
        doc = docx.Document()
        doc.add_heading(draft.title, 0)
        doc.add_paragraph(f"Version: {draft.version}")
        doc.add_paragraph(draft.content)
        buffer = io.BytesIO()
        doc.save(buffer)
        content_bytes = buffer.getvalue()
    else:
        # Fallback PDF generation using reportlab
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        c.drawString(inch, 11 * inch, f"Title: {draft.title}")
        c.drawString(inch, 10.5 * inch, f"Version: {draft.version}")
        textobject = c.beginText(inch, 10 * inch)
        for line in draft.content.split("\n"):
            textobject.textLine(line)
        c.drawText(textobject)
        c.showPage()
        c.save()
        content_bytes = buffer.getvalue()

    s3_key = f"{draft.tenant_id}/{draft.matter_id}/exports/draft_{draft.id}_v{draft.version}.{fmt}"

    try:
        async with storage_service.session.client(
            "s3",
            endpoint_url=storage_service.endpoint_url,
            aws_access_key_id=storage_service.aws_access_key_id,
            aws_secret_access_key=storage_service.aws_secret_access_key,
            config=storage_service.config,
        ) as s3:
            await s3.put_object(
                Bucket=storage_service.bucket_name,
                Key=s3_key,
                Body=content_bytes,
                ContentType=content_type,
            )
        logger.info(f"Successfully exported draft {draft_id} to S3 key {s3_key}")
    except Exception:
        logger.exception("Failed to export draft %s to S3", draft_id)
        raise
